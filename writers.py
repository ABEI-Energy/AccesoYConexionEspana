from math import *

import docx2pdf as d2p
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION

from docx.shared import Cm
from PIL import Image

import numToLet as ntl
from PyPDF2 import PdfFileMerger, PdfFileReader, PdfWriter, PdfReader
import io

#########################################################Doc readers
def doctopdf(docxFile):
    i = 0
    while i < 1:
        try:
            pdfFile = io.BytesIO()
            pdfFile = d2p.convert(docxFile)
            pdfFile.seek(0)
            i +=1
        except AttributeError:
            i +=1
            continue
    return pdfFile

def docDuplicate(docxFile):

    docIn = Document(docxFile)
    docOut = docIn
   
    return docOut

def docWriter(docxFile,docxDict):

    #Headers
    for section in docxFile.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for word in docxDict:
                            if word in paragraph.text:
                                paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                                paragraph.style = docxFile.styles['HeaderStyle'] 


    #Cover
    for paragraph in docxFile.paragraphs:
        for word in docxDict:
            if word in paragraph.text:
                if ((word =="municipioProjC") or (word =="provinciaProjC") or (word =="ccaaProjC") or (word =="dateCoverC") or (word =="versionDocC")) and (paragraph.style.name == "CoverLight"):
                    paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                    paragraph.style = docxFile.styles['CoverLight'] 
                elif ((word=="potPicoC") or (word == "potInstaladaC") or (word == "nombreProyectoC")) and (paragraph.style.name == "CoverBold"):
                    paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                    paragraph.style = docxFile.styles['CoverBold']
                else: pass
    
    # ¿Tiene sentido eliminarlas? 
    # coverKeys = ["municipioProjC","provinciaProjC", "ccaaProjC","dateCoverC", "versionDocC", "potPicoC", "potInstaladaC","nombreProyectoC"]
    # [docxDict.pop(key) for key in coverKeys]

    # priceKeys = {"precioModulo":'', "precioTrafo":'', "precioEstruct":'', "precioInvert":'', "equiposTotal":'', "totalPrecioP":'',  "subtotal":'', "sub5":'', "ind10":'', "totalLP":''}
    # docxDict.update(priceKeys)

    #Table values in general

    for table in docxFile.tables:
        for row in table.rows:
            for cell in row.cells:
                if "COD." in cell.text:
                    priceKeys = {"precioModulo":'', "precioTrafo":'', "precioEstruct":'', "precioInvert":'', "equiposTotal":'', "totalPrecioP":'',  "subtotal":'', "sub5":'', "ind10":'', "totalLP":''}
                    docxDict.update(priceKeys)
                    df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if cell.text:
                                df[i][j] = cell.text
                    df = pd.DataFrame(df)
                    df.columns = df.iloc[0]
                    df = df[1:]
                    df.reset_index(drop = True, inplace = True)
                    df['PRECIO UNITARIO'] = df['PRECIO UNITARIO'].apply(lambda x: x.replace(".", ""))
                    df['PRECIO TOTAL'] = df['PRECIO TOTAL'].apply(lambda x: x.replace(".", ""))
                    df['PRECIO UNITARIO'] = df['PRECIO UNITARIO'].apply(lambda x: x.replace(" €", ""))
                    df['PRECIO TOTAL'] = df['PRECIO TOTAL'].apply(lambda x: x.replace(" €", ""))
                    df['PRECIO UNITARIO'] = df['PRECIO UNITARIO'].apply(lambda x: x.strip())
                    df['PRECIO UNITARIO dummy'] = pd.to_numeric(df['PRECIO UNITARIO'], errors = 'coerce')
                    df['PRECIO TOTAL dummy'] = pd.to_numeric(df['PRECIO TOTAL'], errors = 'coerce')

                    # Obtenemos precios de los distintos componentes
                    precioModulo = float(df['PRECIO UNITARIO dummy'].iloc[1])
                    precioTrafo = float(df['PRECIO UNITARIO dummy'].iloc[2])
                    precioEstructura = float(df['PRECIO UNITARIO dummy'].iloc[3])
                    precioInversor = float(df['PRECIO UNITARIO dummy'].iloc[4])

                    df.loc[df['CANTIDAD']=="numModulos",'PRECIO TOTAL dummy'] = (precioModulo*int(docxDict['numModulos']))
                    docxDict['precioModulo'] = str(precioModulo*int(docxDict['numModulos']))

                    df.loc[df['CANTIDAD']=="numTrafos",'PRECIO TOTAL dummy'] = (precioTrafo*int(docxDict['numTrafos']))
                    docxDict['precioTrafo'] = str(precioTrafo*int(docxDict['numTrafos']))

                    df.loc[df['CANTIDAD']=="numEstructuras",'PRECIO TOTAL dummy'] = (precioEstructura*int(docxDict['numEstructuras']))
                    docxDict['precioEstruct'] = str(precioEstructura*int(docxDict['numEstructuras']))

                    df.loc[df['CANTIDAD']=="numInverter",'PRECIO TOTAL dummy'] = (precioInversor*int(docxDict['numInverter']))
                    docxDict['precioInvert'] = str(precioInversor*int(docxDict['numInverter']))

                    docxDict['equiposTotal'] = sum(df['PRECIO TOTAL dummy'].iloc[1:df[df['DESCRIPCIÓN']=="OBRA CIVIL"].index[0]].values)
                    docxDict['subtotal'] = sum(df['PRECIO TOTAL dummy'].iloc[df[df['DESCRIPCIÓN']=="OBRA CIVIL"].index[0]:df[df['DESCRIPCIÓN']=="ESTUDIO DE SEGURIDAD Y SALUD"].index[0]].values) + docxDict['equiposTotal'] 
                    docxDict['sub5'] = 0.05*docxDict['subtotal']
                    docxDict['ind10'] = 0.1*docxDict['subtotal']
                    docxDict['totalPrecioP'] = str(docxDict['subtotal'] + docxDict['sub5'] + docxDict['ind10'])
                    auxTotalPrecioP = float(docxDict['totalPrecioP'])
                    docxDict['totalPrecioIvaP'] = str(float(docxDict['totalPrecioP'])*1.21)

                    docxDict['subtotal'] = str(docxDict['subtotal'])
                    docxDict['sub5'] = str(docxDict['sub5'])
                    docxDict['ind10'] = str(docxDict['ind10'])

                    docxDict['totalLetraPrecioIva'] = ntl.numero_a_letras(round(float(docxDict['subtotal'])))
                    
                    docxDict['equiposTotal'] = '{:,}'.format(round(float(docxDict['equiposTotal']))).replace(',','.')  
                    docxDict['precioModulo'] = '{:,}'.format(round(float(docxDict['precioModulo']))).replace(',','.')  
                    docxDict['precioTrafo'] = '{:,}'.format(round(float(docxDict['precioTrafo']))).replace(',','.')    
                    docxDict['precioEstruct'] = '{:,}'.format(round(float(docxDict['precioEstruct']))).replace(',','.')
                    docxDict['precioInvert'] = '{:,}'.format(round(float(docxDict['precioInvert']))).replace(',','.')
                    docxDict['subtotal'] = '{:,}'.format(round(float(docxDict['subtotal']))).replace(',','.')
                    docxDict['sub5'] = '{:,}'.format(round(float(docxDict['sub5']))).replace(',','.')
                    docxDict['ind10'] = '{:,}'.format(round(float(docxDict['ind10']))).replace(',','.')
                    docxDict['totalPrecioP'] = '{:,}'.format(round(float(docxDict['totalPrecioP']))).replace(',','.')
                    docxDict['totalPrecioIvaP'] = '{:,}'.format(round(float(docxDict['totalPrecioIvaP']))).replace(',','.')

                elif "PLANTA nombreProyecto" in cell.text:
                    for paragraph in cell.paragraphs:
                        for word in docxDict:
                            if word in paragraph.text:
                                ps1 = paragraph.style
                                paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                                paragraph.style = ps1      
                    df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if cell.text:
                                df[i][j] = cell.text
                    df = pd.DataFrame(df)
                    df.columns = ['ÍNDICE', 'VALORES'] 
                    df['VALORES'] = df['VALORES'].apply(lambda x: x.replace(".", ""))
                    df['VALORES'] = df['VALORES'].apply(lambda x: x.replace(" €", ""))
                    df['VALORES'] = df['VALORES'].apply(lambda x: x.strip())

                    docxDict['costeLinea'] = float(df['VALORES'].iloc[1])
                    docxDict['totalLP'] = auxTotalPrecioP+docxDict['costeLinea']
                    docxDict['totalIvaLP'] = docxDict['totalLP']*1.21
                    docxDict['totalLetraIvaP'] = ntl.numero_a_letras(round(float(docxDict['totalLP'])))

                    docxDict['totalLP'] = '{:,}'.format(round(float(docxDict['totalLP']))).replace(',','.')
                    docxDict['totalIvaLP'] = '{:,}'.format(round(float(docxDict['totalIvaLP']))).replace(',','.')


                for paragraph in cell.paragraphs:
                    previousStyle = paragraph.style
                    for word in docxDict:
                        if word in paragraph.text:
                            paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                            paragraph.style = previousStyle  
                            if word == "FlagReference":
                                pass 
    
        #Resto del documento
    for paragraph in docxFile.paragraphs:
        for word in docxDict:
            if word in paragraph.text:           
                    if word == "FlagReference":
                        paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                        paragraph.style = docxFile.styles['tableCaption']
                    if word == "FlagFigRef":
                        paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                        paragraph.style = docxFile.styles['figureCaption']                         
                    elif word == "trafoLongTab":
                        paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                        paragraph.style = docxFile.styles['bulletDoc']
                    elif ((word=="totalLetraIvaP") or  (word=="totalLetraPrecioIva")):
                        paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                        paragraph.style = docxFile.styles['normalBold']
                    else:
                        previousStyle = paragraph.style
                        paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                        paragraph.style = previousStyle

    flagDocWriter = 1
    print(flagDocWriter)
    return flagDocWriter
                                  
def docTabler(docxFile, dfP, dfT, dfV, dfCT, dfL,dfA):

    #Aquí sólo tocamos las tablas que tengamos que rellenar con dataframes

    for table in docxFile.tables:
        for row in table.rows:
            for cell in row.cells:
                if "plantaFlag" in cell.text:
                    cellStyle = cell.paragraphs[0].style
                    cell.text = ""
                    for i in range(len(dfP)-1):
                        table.add_row()
                        table.style = 'Table Grid'
                    for i in range(dfP.shape[0]):
                        for j in range(dfP.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(dfP.values[i,j])
                            table.cell(i+1,j).paragraphs[0].style = cellStyle
                if "tramoFlag" in cell.text:
                    cellStyle = cell.paragraphs[0].style
                    cell.text = ""
                    for i in range(len(dfT)-1):
                        table.add_row()
                        table.style = 'Table Grid'
                    for i in range(dfT.shape[0]):
                        for j in range(dfT.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(dfT.values[i,j])
                            table.cell(i+1,j).paragraphs[0].style = cellStyle
                if "flagVallado" in cell.text:
                    cellStyle = cell.paragraphs[0].style
                    cell.text = ""
                    for i in range(len(dfV)-1):
                        table.add_row()
                        table.style = 'Table Grid'
                    for i in range(dfV.shape[0]):
                        for j in range(dfV.shape[-1]):
                            table.cell(i+2,j).paragraphs[0].text = str(dfV.values[i,j])
                            table.cell(i+2,j).paragraphs[0].style = cellStyle
                if "flagCTM" in cell.text:
                    cellStyle = cell.paragraphs[0].style
                    cell.text = ""
                    for i in range(len(dfCT)-1):
                        table.add_row()
                        table.style = 'Table Grid'
                    for i in range(dfCT.shape[0]):
                        for j in range(dfCT.shape[-1]):
                            table.cell(i+2,j).paragraphs[0].text = str(dfCT.values[i,j])
                            table.cell(i+2,j).paragraphs[0].style = cellStyle
                if "flagLinea" in cell.text:
                    cellStyle = cell.paragraphs[0].style
                    cell.text = ""
                    for i in range(len(dfL)-1):
                        table.add_row()
                        table.style = 'Table Grid'
                    for i in range(dfL.shape[0]):
                        for j in range(dfL.shape[-1]):
                            table.cell(i+2,j).paragraphs[0].text = str(dfL.values[i,j])
                            table.cell(i+2,j).paragraphs[0].style = cellStyle
                if "flagAcceso" in cell.text:
                    cellStyle = cell.paragraphs[0].style
                    cell.text = ""
                    for i in range(len(dfA)-1):
                        table.add_row()
                        table.style = 'Table Grid'
                    for i in range(dfA.shape[0]):
                        for j in range(dfA.shape[-1]):
                            table.cell(i+2,j).paragraphs[0].text = str(dfA.values[i,j])
                            table.cell(i+2,j).paragraphs[0].style = cellStyle
    print("Documentos docx done (con tablas)") 
    flagDocTabler = 1
    return flagDocTabler

def picWriter(docxFile, docxDict, rootLogos, logoC):

    docxDict['logoC'] = rootLogos + "/" + docxDict['logoC']
    docxDict['logoH'] = rootLogos + "/" + docxDict['logoH']


    for table in docxFile.tables:
        for row in table.rows:
            for cell in row.cells:
                if "logoC" in cell.text:
                    # image = Image.open(docxDict['logoC'])
                    # st.image(image, caption='Uploaded Image', use_column_width=True)
                    # logoC = io.BytesIO()
                    # image.save(logoC, format = 'PNG')
                    # logoC.seek(0)
                    cell.text = ""
                    r = cell.paragraphs[0].add_run()
                    r.add_picture(logoC, width = Cm(7.0), height = Cm(4.5))
                if ("tomaAerea" in cell.text) and (docxDict['tomaAerea'] != "pass"):
                    cell.text = ""
                    r = cell.paragraphs[0].add_run()
                    r.add_picture(docxDict['tomaAerea'], width = Cm(16.55), height = Cm(10.05))
                if ("figuraStruct" in cell.text) and (docxDict['figuraStruct'] != "pass"):
                    cell.text = ""
                    r = cell.paragraphs[0].add_run()
                    r.add_picture(docxDict['figuraStruct'], width = Cm(6.95), height = Cm(5.1)) 
                if ("figuraCronograma" in cell.text) and (docxDict['figuraCronograma'] != "pass"):
                    cell.text = ""
                    r = cell.paragraphs[0].add_run()
                    r.add_picture(docxDict['figuraCronograma'], width = Cm(15.2), height = Cm(14.3))                    
                    
    #Headers
    for section in docxFile.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "logoH" in cell.text:
                        cell.text = ""
                        r = cell.paragraphs[0].add_run()
                        r.add_picture(docxDict['logoH'], width = Cm(2.75), height = Cm(1.75))         
    print("Imagenes insertadas") 

    flagPicWriter = 1
    return flagPicWriter

def pdfMerger(files):
    tmp = io.BytesIO()
    merger = PdfFileMerger()
    for pdf in files:
        merger.append(PdfReader(pdf))
        # pdf = open(pdf,'rb')
        # merger.append(pdf)
    merger.write(tmp)
    tmp.seek(0)
    # return tmp.getvalue()
    return tmp

def pdfInsert(docWord, docPdf,flagPlanos=0):
    pdf_reader = PdfFileReader(docPdf)
    docWord_in = Document(docWord)

    if flagPlanos == 1:
        current_section = docWord_in.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = docWord_in.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        for page_num in range(pdf_reader.numPages):
            # Extract the page from the PDF
            # Convert the page to an image (requires the 'pdf2image' library)

            from pdf2image import convert_from_bytes
            images = convert_from_bytes(docPdf.getvalue(), first_page=page_num+1, last_page=page_num+1)
            image_stream = io.BytesIO()
            images[0].save(image_stream, format='PNG')
            image_stream.seek(0)

            size = images[0].size
            imWidth = float(size[0])
            imHeight = float(size[1])                    
            if imWidth > 23.82:
                imHeight = imHeight*23.82/imWidth
                imWidth = 23.82


            run = docWord_in.paragraphs[-1].add_run()
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.add_picture(image_stream, width = Cm(imWidth), height = Cm(imHeight))
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER


            docWord_out = io.BytesIO()
            docWord_in.save(docWord_out)
            docWord_out.seek(0)
            # Close the image stream
            image_stream.close()

    else:

        # Pasamos el documento ya en bytesIo
         # Iterate over each page in the PDF

        for page_num in range(pdf_reader.numPages):
            # Extract the page from the PDF
            page = pdf_reader.getPage(page_num)

            # Convert the page to an image (requires the 'pdf2image' library)
            # Here, we assume you have the 'pdf2image' library installed
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(docPdf.getvalue(), first_page=page_num+1, last_page=page_num+1)
            image_stream = io.BytesIO()
            images[0].save(image_stream, format='PNG')
            image_stream.seek(0)

            size = images[0].size
            imWidth = float(size[0])
            imHeight = float(size[1])                    
            if imWidth > 17.32:
                imHeight = imHeight*17.32/imWidth
                imWidth = 17.32

            run = docWord_in.paragraphs[-1].add_run()
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.add_picture(image_stream, width = Cm(imWidth), height = Cm(imHeight))
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER


            docWord_out = io.BytesIO()
            docWord_in.save(docWord_out)
            docWord_out.seek(0)
            # Close the image stream
            image_stream.close()


    return docWord_out

# Function to convert DOCX to PDF
def convert_docx_to_pdf(docx_data):


    doc = Document(docx_data)

    # Create a PDF writer
    pdf_writer = PdfWriter()

    # Iterate over each page in the DOCX document
    for page_num, page in enumerate(doc.pages):
        # Extract the content of the page as text
        text = page.text

        # Add the text to the PDF writer
        pdf_writer.addPage(text)

    # Convert the PDF writer to a bytes object
    pdf_output = io.BytesIO()
    pdf_writer.write(pdf_output)
    pdf_output.seek(0)

    return pdf_output



