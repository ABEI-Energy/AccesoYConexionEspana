# import modules
import streamlit as st  # for web app


# Arial check

    
# Abadi check

# dwnld_btn = st.sidebar.download_button(label='Download edited document', data=[edited_doc1, edited_doc2], key='e')

import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image
import io

# def insert_image_in_cell(table, cell_index, image):
#     cell = table.cell(cell_index[0], cell_index[1])
#     cell_paragraph = cell.paragraphs[0]
#     run = cell_paragraph.add_run()
#     image_stream = io.BytesIO()
#     image.save(image_stream, format='PNG')
#     image_stream.seek(0)
#     run.add_picture(image_stream)

def insert_image_in_cell(doc, image):
    table = doc.tables[0]
    cell = table.cell(0,0)
    cell_paragraph = cell.paragraphs[0]
    run = cell_paragraph.add_run()
    image_stream = io.BytesIO()
    image.save(image_stream, format='PNG')
    image_stream.seek(0)
    run.add_picture(image_stream)

def download_document(doc):
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

def main():
    st.title("Image to Word Document")
    uploaded_file = "SOCIEDADES\Logos\ABDA.png"
    # uploaded_file = st.file_uploader("Upload an image", type=['png', 'jpg', 'jpeg'])
    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption='Uploaded Image', use_column_width=True)

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        insert_image_in_cell(table, (0, 0), image)

        bio = io.BytesIO()
        doc.save(bio)
        st.download_button(label="Descargar documento", data=bio.getvalue(), file_name="doc.docx", mime = "docx")

if __name__ == '__main__':
    main()
