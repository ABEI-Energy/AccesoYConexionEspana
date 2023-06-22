# '''# import modules
# from docx import Document  # for accessing the document by python-docx
# import streamlit as st  # for web app


# # Arial check
# def func1():
#     global edited_doc1
#     # access Word document
#     my_file = Document("MODELOS\Modelo anteproyecto AFL.docx")
#     for paragraph in my_file.paragraphs:
#         st.write(paragraph.text)
    
# # Abadi check
# def func2():
#     global edited_doc2
#     # access Word document
#     my_file = Document(uploaded_file)
#     FONT = 'Abadi'  # state the specified font
#     for paragraphs in my_file.paragraphs:
#         if 'Normal' == paragraphs.style.name:
#             for run in paragraphs.runs:
#                 if run.font.name == FONT:
#                     # underline text written in Abadi font
#                     run.underline = True
#                     # save changes to file
#                     edited_doc2 = my_file.save("Edited_Document2.docx")
#     return edited_doc2


# # function to run selected programs(s)
# def run_program():
#     if cb1: func1()
#     if cb2: func2()


# # configure sidebar text and widgets
# uploaded_file = st.sidebar.file_uploader("Select Word document", type='.docx', key='a')
# cb1 = st.sidebar.checkbox('Arial check', key='b')
# cb2 = st.sidebar.checkbox('Abadi check', key='c')
# run_btn = st.sidebar.button('Run', on_click=run_program, key='d')
# # dwnld_btn = st.sidebar.download_button(label='Download edited document', data=[edited_doc1, edited_doc2], key='e')'''

# '''import streamlit as st
# from docx import Document
# from PIL import Image
# import io

# def insert_image(doc, image):
#     image_stream = io.BytesIO()
#     image.save(image_stream, format='PNG')
#     image_stream.seek(0)
#     doc.add_picture(image_stream)

# def download_document(doc):
#     doc_stream = io.BytesIO()
#     doc.save(doc_stream)
#     doc_stream.seek(0)
#     return doc_stream

# def main():
#     st.title("Image to Word Document")

#     uploaded_file = "SOCIEDADES\Logos\ABDA.png"

#     if uploaded_file is not None:
#         image = Image.open(uploaded_file)
#         st.image(image, caption='Uploaded Image', use_column_width=True)

#         doc = Document()
#         insert_image(doc, image)


#         bio = io.BytesIO()
#         doc.save(bio)
#         # docPdf = wt.doctopdf(doc_modelo)
#         st.download_button(label="Descargar documento", data=bio.getvalue(), file_name="doc.docx", mime = "docx")

# if __name__ == '__main__':
#     main()'''


# import streamlit as st
# from docx import Document
# from docx.shared import Inches
# from PIL import Image
# import io

# def insert_image_in_cell(table, cell_index, image):
#     cell = table.cell(cell_index[0], cell_index[1])
#     cell_paragraph = cell.paragraphs[0]
#     run = cell_paragraph.add_run()
#     image_stream = io.BytesIO()
#     image.save(image_stream, format='PNG')
#     image_stream.seek(0)
#     run.add_picture(image_stream)

# def download_document(doc):
#     doc_stream = io.BytesIO()
#     doc.save(doc_stream)
#     doc_stream.seek(0)
#     return doc_stream

# def main():
#     st.title("Image to Word Document")
#     uploaded_file = "SOCIEDADES\Logos\ABDA.png"
#     # uploaded_file = st.file_uploader("Upload an image", type=['png', 'jpg', 'jpeg'])
#     if uploaded_file is not None:
#         image = Image.open(uploaded_file)
#         st.image(image, caption='Uploaded Image', use_column_width=True)

#         doc = Document()
#         table = doc.add_table(rows=1, cols=1)
#         insert_image_in_cell(table, (0, 0), image)

#         bio = io.BytesIO()
#         doc.save(bio)
#         st.download_button(label="Descargar documento", data=bio.getvalue(), file_name="doc.docx", mime = "docx")

# if __name__ == '__main__':
#     main()


import streamlit as st
import pandas as pd
import streamlit_toggle as tog
import io
import readers as rd
import coordinates as cd
from aereas import * 
import readers as rd
import writers as wt
import post as ps
from docx.shared import Cm 
import locale as lc
# from docx import Document  # for accessing the document by python-docx
import datetime as dt
import numToLet as ntl


# Sociedades
rootSociedades = 'SOCIEDADES/Sociedades España.csv'
rootEstructuras = "DATASHEETS/Estructuras"
rootWord = "MODELOS"
rootLogos = "SOCIEDADES/Logos"
csvTensiones = 'DATASHEETS/Tensiones maximas.csv'
csvTrafos = 'DATASHEETS/Trafos/Datasheet Trafos.csv'
csvModulos = 'DATASHEETS/Módulos/Datasheet Módulos.csv'
csvInverters = 'DATASHEETS/Inversores/Datasheet Inversores.csv'
csvCeldas = 'DATASHEETS/CeldasMT/CeldasMT.csv'


rootSampleFiles = "SAMPLEFILES/Ejemplo archivos PSFV XXXX.zip"

dfSociedad = pd.read_csv(rootSociedades)
dfTensiones = pd.read_csv(csvTensiones)
dfTrafos = pd.read_csv(csvTrafos)
dfModulos = pd.read_csv(csvModulos)
dfInverters = pd.read_csv(csvInverters)
dfCeldas = pd.read_csv(csvCeldas)


# Geografía
rootProvincias = "GEOGRAFÍA/Provincias.csv"
rootMunicipios =  "GEOGRAFÍA/Municipios.csv"
dfProvincias = pd.read_csv(rootProvincias)
dfMunicipios = pd.read_csv(rootMunicipios)